# -*- coding: utf-8 -*-
"""生成《项目进度表》与《7 月上线总控表》两个 docx，复用任务分配文档的样式。"""
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

BASE = Path(r"E:\小程序开发\docte")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F7F9FC"
RISK_FILL = "FFF2CC"
PASS_FILL = "E2F0D9"
BORDER = "B7C9DD"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_cm):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, name="Calibri", size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def fill_cell(cell, text, bold=False, size=9.2, color=None, align=None, fill=None):
    cell.text = ""
    para = cell.paragraphs[0]
    if align is not None:
        para.alignment = align
    para.paragraph_format.space_after = Pt(1)
    para.paragraph_format.line_spacing = 1.12
    for index, line in enumerate(str(text).split("\n")):
        if index:
            para.add_run().add_break()
        run = para.add_run(line)
        set_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)
    set_cell_border(cell)
    if fill:
        set_cell_shading(cell, fill)


# 状态文字 -> 行底色
STATUS_FILL = {
    "✅": PASS_FILL,
    "🟡": "FFF8E1",
    "⛔": "F8D7DA",
    "❌": "FBE4E6",
    "⚠️": RISK_FILL,
}


def status_fill_for(text):
    for key, fill in STATUS_FILL.items():
        if key in str(text):
            return fill
    return None


def add_table(doc, headers, rows, widths_cm, font_size=8.2, status_col=None):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths_cm)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        fill_cell(hdr.cells[i], h, bold=True, size=8.6, color="000000",
                  align=WD_ALIGN_PARAGRAPH.CENTER, fill=HEADER_FILL)
    for row in rows:
        cells = table.add_row().cells
        row_fill = status_fill_for(row[status_col]) if status_col is not None else None
        for i, value in enumerate(row):
            fill_cell(cells[i], value, size=font_size)
            if row_fill:
                set_cell_shading(cells[i], row_fill)
            elif i == 0:
                set_cell_shading(cells[i], LIGHT_FILL)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    para = doc.add_heading(level=level)
    para.paragraph_format.space_before = Pt(10 if level > 1 else 14)
    para.paragraph_format.space_after = Pt(5)
    run = para.add_run(text)
    set_font(run, size=16 if level == 1 else 13 if level == 2 else 11.5,
             bold=True, color=BLUE if level < 3 else DARK_BLUE)
    return para


def add_note_table(doc, title, body, fill):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [16.2])
    cell = table.rows[0].cells[0]
    fill_cell(cell, f"{title}\n{body}", bold=False, size=9.2, fill=fill)
    first_run = cell.paragraphs[0].runs[0]
    first_run.bold = True
    first_run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    doc.add_paragraph()


def new_doc(title, subtitle, note=None):
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    for style_name in ("Normal", "Body Text"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.18
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run(title), size=8.5, color="666666")

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_after = Pt(4)
    set_font(t.add_run(title), size=20, bold=True, color=DARK_BLUE)
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_after = Pt(10)
    set_font(s.add_run(subtitle), size=10.5, color="555555")
    if note:
        add_note_table(doc, note[0], note[1], note[2])
    return doc


# ---------------- 进度表 ----------------
def build_progress():
    doc = new_doc(
        "牙医仪器检修小程序 · 项目进度表",
        "统计时间：2026-06-22 ｜ 目标：2026 年 7 月上线",
        ("总体判断", "小程序与后台代码功能已基本完成（16/16 主流程 + 合规整改），上线卡点为：①微信医疗器械类目审核（阻断）；②正式环境部署与微信支付/订阅/域名/索引等控制台配置；③真机与全流程验收尚未开始。", PASS_FILL),
    )

    add_heading(doc, "一、总览", 1)
    add_table(doc, ["维度", "完成度", "说明"], [
        ["小程序客户端代码功能", "✅ ~95%（16/16）", "仅余号码/资质等内容确认"],
        ["后端云函数代码", "✅ 已完成", "7 个云函数齐备，合规整改已写完"],
        ["后端部署（正式环境）", "❌ 未完成", "仅部署到开发云空间，5 个函数需重部署"],
        ["数据库索引", "❌ 未完成", "清单已写，控制台未创建"],
        ["微信平台/支付/订阅配置", "❌ 未完成", "域名、支付、模板 ID 均未配"],
        ["真机/全流程验收", "❌ 未开始", "排期 06-28 起"],
        ["类目审核（阻断项）", "⛔ 未通过", "医疗器械类目资质审核未过"],
    ], [4.0, 4.2, 8.0], font_size=8.6, status_col=1)

    add_heading(doc, "二、小程序客户端代码功能（16 / 16）", 1)
    feats = [
        ["1", "手机验证码登录 + token + 401 跳登录", "✅", "pages/login、api/auth.js、cloudHelpers.js"],
        ["2", "拒绝授权可恢复路径", "✅", "login/index.vue、PrivacyConsent.vue"],
        ["3", "报修建单（选设备 SN/回寄地址自动带入）", "✅", "api/repair.js、createOrder"],
        ["4", "报修上传图片", "✅", "uploadRepairImage、uploadToCloud"],
        ["5", "订单进度/状态（列表+详情）", "✅", "getRepairList/Detail、getOrderList"],
        ["6", "查看/确认/拒绝报价", "✅", "confirm/rejectRepairQuote、rejectQuote"],
        ["7", "微信支付 requestPayment", "✅", "index.vue L2830、createWechatPayPayment"],
        ["8", "对公转账/付款凭证上传", "✅", "uploadRepairPaymentProof"],
        ["9", "发票申请", "✅", "api/content.js applyInvoice"],
        ["10", "确认收货 confirmReceipt", "✅", "confirmRepairReceipt 云函数"],
        ["11", "评价/回访（不满意转投诉）", "✅", "reviewOrder、submitOrderReview"],
        ["12", "投诉/建议提交 + 查看列表", "✅", "addComplaint/getComplaintList"],
        ["13", "订阅消息授权（全场景）", "✅", "requestStatusSubscription、sendOrderSubscription"],
        ["14", "设备档案页 listMyDevices", "✅", "getMyDevices、listMyDevices"],
        ["15", "隐私授权弹窗 + onNeedPrivacyAuthorization", "✅", "PrivacyConsent.vue"],
        ["16", "账号注销（软删除+脱敏）", "✅", "cancelAccount 云函数"],
    ]
    add_table(doc, ["#", "功能", "状态", "证据"], feats, [0.9, 7.2, 1.6, 6.5], font_size=8.0, status_col=2)

    add_heading(doc, "三、最低上线门槛（12 项 · 达成 0/12）", 1)
    gates = [
        ["1", "正式 AppID 真机登录成功", "✅", "❌ 真机未验收", "❌"],
        ["2", "正式库可建工单，后台可见", "✅", "❌ 未连正式环境", "❌"],
        ["3", "后台改状态，小程序进度同步", "✅", "❌ 未验收", "❌"],
        ["4", "后台发报价，小程序可确认", "✅", "❌ 未验收", "❌"],
        ["5", "微信支付或对公凭证至少一条可用", "✅", "❌ WX_PAY_* 未配置", "❌"],
        ["6", "发票申请与后台开票状态同步", "✅", "❌ 未验收", "❌"],
        ["7", "回寄物流小程序可见", "✅", "❌ 未验收", "❌"],
        ["8", "订阅消息成功/失败/拒绝可追踪", "✅", "❌ 模板 ID 未配", "❌"],
        ["9", "PC 后台正式域名 HTTPS 可访问", "—", "❌ 未部署", "❌"],
        ["10", "不依赖 localhost/mock/开发云空间", "—", "❌ 当前连开发云", "❌"],
        ["11", "数据库唯一+关键复合索引已创建", "✅清单", "❌ 控制台未创建", "❌"],
        ["12", "客户可见资料全部正式版本", "—", "⚠️ 热线/资质待确认", "⚠️"],
    ]
    add_table(doc, ["#", "门槛", "代码", "部署/配置/验收", "综合"],
              gates, [0.9, 7.6, 1.8, 4.4, 1.5], font_size=8.0, status_col=4)

    add_heading(doc, "四、上线里程碑", 1)
    add_table(doc, ["阶段", "排期", "状态"], [
        ["正式资料锁定", "06-22 ~ 06-24", "🟡 进行中"],
        ["正式环境部署", "06-25 ~ 06-27", "⬜ 未开始"],
        ["第一轮全流程验收", "06-28 ~ 06-30", "⬜ 未开始"],
        ["最终回归与提审准备", "07-01 ~ 07-03", "⬜ 未开始"],
        ["微信审核与发布", "07-04 起", "⬜ 未开始"],
    ], [5.0, 5.0, 6.0], font_size=8.6, status_col=2)

    add_note_table(doc, "一句话总结",
                   "开发已基本完成，上线在配置和审核。代码侧无重大缺口，剩下的是“配置 + 审核 + 验收”三类工作。",
                   RISK_FILL)
    out = BASE / "项目进度表_2026-06-22.docx"
    doc.save(out)
    print(out)


# ---------------- 上线总控表 ----------------
def build_control():
    doc = new_doc(
        "7 月上线总控表",
        "负责人：吴彤彤 ｜ 建立：2026-06-22 ｜ 目标：2026 年 7 月上线",
        ("使用说明", "覆盖 6 人全部任务，每天至少更新一次；阻塞项必须有下一步动作。状态：✅已完成 / 🟡进行中 / ⬜未开始 / ⛔阻断。", PASS_FILL),
    )
    H = ["序号", "任务", "负责人", "依赖", "状态", "阻塞原因 / 下一步动作", "预计完成"]
    W = [1.2, 4.6, 1.8, 2.6, 1.2, 4.4, 2.0]

    def section(title, rows):
        add_heading(doc, title, 2)
        add_table(doc, H, rows, W, font_size=7.6, status_col=4)

    add_heading(doc, "一、阻断项（最高优先）", 1)
    add_table(doc, H, [
        ["B1", "微信医疗器械类目审核通过", "吴彤彤", "客户资质证照", "⛔",
         "类目审核未通过；补齐《医疗器械生产许可证》《互联网信息服务备案凭证》后在「类目管理」重新提交", "待定（卡发布）"],
    ], W, font_size=7.6, status_col=4)

    add_heading(doc, "二、按人分工", 1)
    section("客户资料 · 吴彤彤", [
        ["A1", "建立并每日维护本总控表", "吴彤彤", "无", "✅", "本表已建立，进入每日更新", "2026-06-22"],
        ["A2", "收齐正式资料（AppID/云空间/商户号/PC 域名/客服/Logo/介绍/政策）", "吴彤彤", "客户提供", "⬜", "向客户索取并逐项登记状态", "2026-06-24"],
        ["A3", "确认权威客服热线号", "吴彤彤", "客户确认", "🟡", "暂用 0757-85775667，待客户确认", "2026-06-24"],
        ["A4", "组织每日进度同步", "吴彤彤", "各人反馈", "🟡", "已启动，跟踪支付/订阅/URL/索引/真机", "2026-07-03"],
    ])
    section("微信平台/支付/订阅 · 陈莎", [
        ["C1", "配置合法域名（request/upload/download）", "陈莎", "正式域名", "⬜", "依赖 D4 URL 化、F3 正式域名", "2026-06-24"],
        ["C2", "正式 AppID 手机号授权 + openid", "陈莎", "E2 体验版", "⬜", "代码已实现，待真机验证", "2026-06-25"],
        ["C3", "申请订阅模板 ID（报修/签收/报价/付款/发货/完成/回访）", "陈莎", "平台审核", "⬜", "场景代码已就绪，待申请回填", "2026-06-26"],
        ["C4", "微信支付商户号/证书/APIv3/回调", "陈莎", "客户商户资料", "⬜", "退款代码已就绪，待 WX_PAY_*", "2026-06-26"],
        ["C5", "小额真实支付 + 订阅消息测试", "陈莎", "D/E 联调", "⬜", "依赖支付配置与正式环境", "2026-06-28"],
    ])
    section("uniCloud/数据库/云函数 · 苏凤冰", [
        ["D1", "创建/接入客户正式云空间", "苏凤冰", "客户账号权限", "⬜", "当前开发云 env-00jy6bcqqsjw，待切正式", "2026-06-25"],
        ["D2", "部署 7 个云函数（含整改重部署 5 个）", "苏凤冰", "代码冻结", "🟡", "代码就绪；曾部署开发云，需重部署正式", "2026-06-26"],
        ["D3", "配置正式环境变量（登录/支付/订阅）", "苏凤冰", "陈莎参数", "⬜", "依赖 C3/C4", "2026-06-27"],
        ["D4", "PC 后台云函数 URL 化并交付", "苏凤冰", "D2", "⬜", "交蔡珠琪配置", "2026-06-27"],
        ["D5", "创建集合/schema/索引（order_no 唯一、sn 唯一前清洗）", "苏凤冰", "正式云空间", "⬜", "清单已写，控制台未创建", "2026-06-27"],
    ])
    section("小程序端/真机 · 郑佳华", [
        ["E1", "manifest/project 用正式 AppID 且无密钥", "郑佳华", "C 正式 AppID", "🟡", "当前 wx25289fbe4a3bf011，待确认正式主体", "2026-06-26"],
        ["E2", "构建体验版（二维码+版本号）", "郑佳华", "正式云函数", "🟡", "已生产构建，待连正式环境", "2026-06-27"],
        ["E3", "真机跑主流程（登录→…→评价）", "郑佳华", "后台配合", "⬜", "16 项功能已实现，待真机验收", "2026-06-29"],
        ["E4", "验证订阅授权（允许/拒绝/失败不阻塞）", "郑佳华", "C3", "⬜", "依赖模板配置", "2026-06-29"],
        ["E5", "整理小程序问题清单并复测", "郑佳华", "第一轮验收", "⬜", "—", "2026-06-30"],
    ])
    section("PC 后台部署/验收 · 蔡珠琪", [
        ["F1", "配置 PC 正式 VITE_* URL", "蔡珠琪", "D4", "⬜", "缺配置会报错（已加兜底）", "2026-06-27"],
        ["F2", "生产构建 + 检查命令", "蔡珠琪", "正式 URL", "⬜", "检查脚本已就绪", "2026-06-28"],
        ["F3", "部署正式域名 + HTTPS", "蔡珠琪", "部署权限", "⬜", "—", "2026-06-28"],
        ["F4", "验收后台核心业务（工单/报价/库存/结算/发票/物流/审计/统计）", "蔡珠琪", "云函数+DB", "⬜", "功能已就绪，待验收", "2026-06-30"],
        ["F5", "复核角色权限与敏感按钮（成本/退款脱敏）", "蔡珠琪", "账号角色", "🟡", "RBAC 脱敏已实现，待账号验收", "2026-06-30"],
    ])
    section("资料核对/全流程验收 · 钟浩霆", [
        ["G1", "核对客户可见资料", "钟浩霆", "A2", "⬜", "依赖资料收齐", "2026-06-26"],
        ["G2", "真实业务跑完整闭环", "钟浩霆", "前后台联调", "⬜", "—", "2026-06-29"],
        ["G3", "建立问题清单", "钟浩霆", "验收产生", "⬜", "—", "2026-06-30"],
        ["G4", "最终复验（阻塞清零）", "钟浩霆", "修复完成", "⬜", "—", "2026-07-02"],
        ["G5", "给出最终上线建议", "钟浩霆", "最终复验", "⬜", "—", "2026-07-02"],
    ])

    add_note_table(doc, "上线判定规则",
                   "任一最低上线门槛未过原则上不得提审；有条件上线须由吴彤彤在《最终上线确认单》写明风险、兜底方案、负责人、预计修复时间。",
                   RISK_FILL)
    out = BASE / "7月上线总控表_2026-06-22.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    build_progress()
    build_control()
