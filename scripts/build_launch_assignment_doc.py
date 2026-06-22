from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = Path(r"E:\小程序开发\docte\微信小程序7月上线任务分配与交付标准.docx")


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F7F9FC"
RISK_FILL = "FFF2CC"
PASS_FILL = "E2F0D9"
BORDER = "B7C9DD"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_cm):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, name="Calibri", size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(cell_or_doc, text="", bold=False, size=9.5, color=None, align=None):
    para = cell_or_doc.add_paragraph()
    if align is not None:
        para.alignment = align
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.line_spacing = 1.15
    run = para.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return para


def fill_cell(cell, text, bold=False, size=9.2, color=None, align=None):
    cell.text = ""
    para = cell.paragraphs[0]
    if align is not None:
        para.alignment = align
    para.paragraph_format.space_after = Pt(1)
    para.paragraph_format.line_spacing = 1.12
    for index, line in enumerate(str(text).split("\n")):
        if index:
            para.add_run().add_break()
        run = para.add_run(line)
        set_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)
    set_cell_border(cell)


def add_table(doc, headers, rows, widths_cm, font_size=8.2):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths_cm)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        fill_cell(hdr.cells[i], h, bold=True, size=8.6, color="000000", align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(hdr.cells[i], HEADER_FILL)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            fill_cell(cells[i], value, size=font_size)
            if i == 0:
                set_cell_shading(cells[i], LIGHT_FILL)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    para = doc.add_heading(level=level)
    para.paragraph_format.space_before = Pt(10 if level > 1 else 14)
    para.paragraph_format.space_after = Pt(5)
    run = para.add_run(text)
    set_font(run, size=16 if level == 1 else 13 if level == 2 else 11.5, bold=True, color=BLUE if level < 3 else DARK_BLUE)
    return para


def add_note_table(doc, title, body, fill):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [16.2])
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    fill_cell(cell, f"{title}\n{body}", bold=False, size=9.2)
    first_run = cell.paragraphs[0].runs[0]
    first_run.bold = True
    first_run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    doc.add_paragraph()


def apply_document_style(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    for style_name in ("Normal", "Body Text"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.18

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("微信小程序 7 月上线任务分配与交付标准")
    set_font(run, size=8.5, color="666666")


def build():
    doc = Document()
    apply_document_style(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("微信小程序 7 月上线任务分配与交付标准")
    set_font(run, size=20, bold=True, color=DARK_BLUE)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(10)
    run = sub.add_run("版本：2026-06-21 整理版｜目标：客户 2026 年 7 月上线")
    set_font(run, size=10.5, color="555555")

    add_note_table(
        doc,
        "当前判断",
        "代码侧关键上线风险已完成首轮处理：库存扣减并发、退款并发、PC 后台开发云空间兜底、订单 schema、索引清单和后台核心入口已处理。后续重心转为正式环境、微信平台、支付、订阅消息、真机闭环、客户资料和上线提审。",
        PASS_FILL,
    )

    add_heading(doc, "一、上线总里程碑", 1)
    add_table(
        doc,
        ["阶段", "时间", "目标", "完成标准", "负责人"],
        [
            ["正式资料锁定", "2026-06-22 至 2026-06-24", "确认客户正式主体、账号、资料、政策文档和外部依赖", "资料状态表无空项；敏感信息只记录保管人和配置状态，不进入代码仓库", "吴彤彤"],
            ["正式环境部署", "2026-06-25 至 2026-06-27", "正式 uniCloud、云函数、URL 化、PC 后台、小程序 AppID、微信支付完成联调", "小程序和后台均连接正式环境；检查命令通过；不依赖 localhost、mock、开发云空间", "苏凤冰、蔡珠琪、郑佳华"],
            ["第一轮全流程验收", "2026-06-28 至 2026-06-30", "从客户报修到后台处理、报价、支付、物流、评价、发票跑通", "阻塞问题清零；非阻塞问题有负责人和修复时间", "钟浩霆、吴彤彤"],
            ["最终回归与提审准备", "2026-07-01 至 2026-07-03", "体验版确认、资料确认、上线门槛逐项打勾", "形成最终上线确认单；客户确认可提交审核", "吴彤彤"],
            ["微信审核与发布", "2026-07-04 起", "按客户确认时间提交微信审核并发布", "微信审核通过；发布后首日监控无阻塞故障", "吴彤彤统筹，全员待命"],
        ],
        [2.4, 3.0, 4.2, 5.2, 1.4],
    )

    add_heading(doc, "二、人员职责总表", 1)
    add_table(
        doc,
        ["人员", "角色定位", "本轮重点", "关键交付物", "最终截止"],
        [
            ["吴彤彤", "上线总控 / 客户资料 / 风险推进", "资料收口、排期跟进、阻塞问题推进、上线判断", "上线总控表、正式资料确认表、阻塞问题清单、最终上线确认单", "2026-07-03"],
            ["陈莎", "微信平台 / 支付 / 订阅消息", "微信合法域名、手机号授权、订阅模板、支付商户和回调", "合法域名截图、模板 ID 清单、支付配置确认表、小额支付记录", "2026-06-28"],
            ["苏凤冰", "uniCloud / 数据库 / 云函数", "正式云空间、云函数部署、URL 化、环境变量、数据库索引", "云空间 ID、部署清单、URL 化地址、环境变量确认表、索引确认表", "2026-06-27"],
            ["郑佳华", "小程序端构建 / 真机闭环", "正式 AppID 构建、体验版、真机跑完整客户流程", "体验版二维码、真机测试记录、小程序问题清单", "2026-06-30"],
            ["蔡珠琪", "PC 后台部署 / 后台业务验收", "PC 正式 URL、生产构建、后台核心业务验收", "PC 正式地址、构建结果、检查命令结果、后台验收记录", "2026-06-30"],
            ["钟浩霆", "资料核对 / 全流程验收", "客户可见内容核对、真实业务流程验收、问题清单", "资料确认表、全流程验收记录、问题清单、最终验收建议", "2026-07-02"],
        ],
        [1.8, 3.0, 4.1, 4.6, 2.1],
    )

    add_heading(doc, "三、个人任务与交付标准", 1)
    people = [
        (
            "吴彤彤",
            "上线总控 / 客户资料 / 风险推进",
            [
                ["1", "建立《7 月上线总控表》，字段包含任务、负责人、依赖、状态、阻塞原因、预计完成时间。", "表内覆盖 6 人全部任务；每天至少更新一次；阻塞项必须有下一步动作。", "上线总控表", "2026-06-22", "无"],
                ["2", "收齐正式资料：AppID、云空间、商户号、PC 域名、客服资料、Logo、公司介绍、政策文档。", "资料状态不得出现“未知”；敏感信息只记录保管人和配置状态。", "正式资料确认表", "2026-06-24", "客户提供正式资料"],
                ["3", "组织每日进度同步，重点跟踪微信支付、订阅消息、URL 化、索引、真机测试。", "每日形成 1 条进度记录；阻塞项当天分配责任人。", "每日进度记录", "2026-07-03", "2-6 号反馈"],
                ["4", "组织两轮验收并输出上线判断。", "第一轮阻塞问题清零；最终确认单由客户或项目负责人确认。", "最终上线确认单", "2026-07-03", "全员完成交付"],
            ],
        ),
        (
            "陈莎",
            "微信平台 / 支付 / 订阅消息",
            [
                ["1", "配置微信公众平台合法域名：request、uploadFile、downloadFile。", "正式域名全部通过微信平台校验；截图可追溯。", "合法域名截图", "2026-06-24", "PC/云函数正式域名"],
                ["2", "确认正式 AppID 手机号授权和 openid 获取。", "真机可完成登录；拒绝授权时页面有可恢复路径。", "登录授权测试记录", "2026-06-25", "郑佳华提供体验版"],
                ["3", "申请并整理订阅消息模板 ID。", "至少覆盖报修提交、签收、报价发布、付款确认、回寄发货、完成提醒。", "模板 ID 清单", "2026-06-26", "微信平台审核"],
                ["4", "确认正式微信支付商户号、证书、APIv3 密钥和回调 URL。", "支付参数已交给苏凤冰配置；回调指向正式云函数。", "支付配置确认表", "2026-06-26", "客户商户资料"],
                ["5", "配合完成一笔小额真实支付和订阅消息测试。", "支付成功、回调成功、后台可见付款状态；消息成功/失败可追踪。", "支付与订阅测试记录", "2026-06-28", "苏凤冰、郑佳华联调"],
            ],
        ),
        (
            "苏凤冰",
            "uniCloud / 数据库 / 云函数",
            [
                ["1", "创建或接入客户正式 uniCloud 云空间。", "云空间归属客户正式主体；记录云空间 ID。", "正式云空间 ID", "2026-06-25", "客户提供账号权限"],
                ["2", "部署客户端和后台云函数。", "client-user、client-order、client-public、admin-sys、admin-kb、admin-order、admin-customer 部署完成。", "云函数部署清单", "2026-06-26", "代码冻结版本"],
                ["3", "配置正式环境变量：微信登录、支付、订阅消息等。", "配置项齐全；敏感值不写入仓库；调用测试无缺参错误。", "环境变量确认表", "2026-06-27", "陈莎提供参数"],
                ["4", "开启 PC 后台云函数 URL 化并交付地址。", "URL 可访问；无开发云空间地址；交给蔡珠琪配置 PC 后台。", "URL 化地址清单", "2026-06-27", "云函数部署完成"],
                ["5", "创建正式数据库集合、schema 和索引。", "按 INDEX_TASK.md 创建索引；SN 唯一索引前完成重复/空值检查。", "数据库与索引确认表", "2026-06-27", "正式云空间可用"],
            ],
        ),
        (
            "郑佳华",
            "小程序端构建 / 真机闭环",
            [
                ["1", "确认 manifest.json 和 project.config.json 使用正式 AppID。", "构建包显示正式 AppID；前端代码不包含 AppSecret 或支付密钥。", "配置检查记录", "2026-06-26", "陈莎提供正式 AppID"],
                ["2", "使用微信开发者工具构建体验版。", "体验版二维码可扫码；版本号、构建时间已记录。", "体验版二维码/版本号", "2026-06-27", "正式云函数可用"],
                ["3", "真机跑客户主流程。", "登录、报修、上传图片、查进度、看报价、确认报价、支付/凭证、发票、收货、评价均可完成。", "真机测试记录", "2026-06-29", "后台流程配合"],
                ["4", "验证订阅消息授权体验。", "允许、拒绝、发送失败均不阻塞主流程；结果可追踪。", "订阅授权测试记录", "2026-06-29", "陈莎模板配置"],
                ["5", "整理小程序问题清单并复测。", "每个问题有截图、机型、微信版本、复现步骤、负责人、优先级。", "小程序问题清单", "2026-06-30", "第一轮验收反馈"],
            ],
        ),
        (
            "蔡珠琪",
            "PC 后台部署 / 后台业务验收",
            [
                ["1", "配置 PC 后台正式 VITE_* URL。", "不再指向开发云空间；缺配置时构建或运行会报错。", "环境配置记录", "2026-06-27", "苏凤冰提供 URL"],
                ["2", "执行生产构建和检查命令。", "npm run build、check:errors、check:urls、check:subscription 通过或有明确处理记录。", "构建与检查结果", "2026-06-28", "正式 URL 可用"],
                ["3", "部署 PC 后台正式域名和 HTTPS。", "正式域名可访问；HTTPS 正常；登录页可打开。", "PC 正式访问地址", "2026-06-28", "部署平台权限"],
                ["4", "验收后台核心业务。", "登录、工单、报价、库存、结算、发票、物流、待办、审计日志、运营统计可操作。", "后台验收记录", "2026-06-30", "云函数和数据库可用"],
                ["5", "复核角色权限和敏感按钮显示。", "无权限角色看不到成本、退款等敏感入口；后端仍保持拦截。", "权限验收记录", "2026-06-30", "账号角色准备"],
            ],
        ),
        (
            "钟浩霆",
            "资料核对 / 全流程验收",
            [
                ["1", "核对客户可见资料。", "客服电话、微信、二维码、寄修地址、Logo、公司介绍、政策文档均为正式版本。", "资料确认表", "2026-06-26", "吴彤彤收集资料"],
                ["2", "按真实业务跑完整闭环。", "报修、签收、报价、确认、支付/凭证、核销、回寄、收货、评价、发票申请全链路完成。", "全流程验收记录", "2026-06-29", "前后台联调完成"],
                ["3", "建立问题清单。", "每个问题必须包含描述、截图、复现步骤、负责人、优先级、是否阻塞上线。", "问题清单", "2026-06-30", "验收过程中产生"],
                ["4", "完成最终复验。", "阻塞项全部关闭；非阻塞项有上线后处理计划。", "最终复验记录", "2026-07-02", "修复完成"],
                ["5", "给出最终上线建议。", "明确可上线、有条件上线或暂缓上线，并说明理由。", "最终验收建议", "2026-07-02", "最终复验完成"],
            ],
        ),
    ]

    for name, role, rows in people:
        add_heading(doc, f"{name}：{role}", 2)
        add_table(
            doc,
            ["序号", "任务", "交付标准", "交付物", "截止时间", "依赖"],
            rows,
            [0.9, 4.2, 4.8, 2.5, 1.8, 2.0],
            font_size=7.6,
        )

    add_heading(doc, "四、最低上线门槛", 1)
    gate_rows = [
        ["1", "正式 AppID 小程序真机登录成功", "郑佳华", "真机测试记录"],
        ["2", "正式库可创建报修工单，后台可见", "苏凤冰、蔡珠琪", "工单截图"],
        ["3", "后台状态变更后，小程序进度同步", "郑佳华、蔡珠琪", "前后台截图"],
        ["4", "后台发布报价，小程序可确认", "蔡珠琪、郑佳华", "报价确认记录"],
        ["5", "微信支付或对公凭证至少一条付款路径完整可用", "陈莎、郑佳华、蔡珠琪", "支付/核销记录"],
        ["6", "发票申请和后台开票状态同步", "蔡珠琪、钟浩霆", "发票流程截图"],
        ["7", "回寄物流在小程序可见", "蔡珠琪、郑佳华", "物流展示截图"],
        ["8", "订阅消息成功、失败、拒绝授权均可追踪且不阻塞流程", "陈莎、钟浩霆", "订阅测试记录"],
        ["9", "PC 后台正式域名 HTTPS 可访问", "蔡珠琪", "正式访问地址"],
        ["10", "正式环境不依赖 localhost、mock、开发云空间、测试账号", "苏凤冰、蔡珠琪", "环境确认记录"],
        ["11", "数据库唯一索引和关键复合索引已创建", "苏凤冰", "索引确认表"],
        ["12", "客户可见资料全部为正式版本", "钟浩霆、吴彤彤", "资料确认表"],
    ]
    add_table(doc, ["序号", "上线门槛", "负责人", "验收证据"], gate_rows, [0.9, 8.0, 3.2, 4.1], font_size=8.2)

    add_heading(doc, "五、每日跟进模板", 1)
    add_table(
        doc,
        ["日期", "负责人", "今日完成", "当前阻塞", "明日计划", "是否影响上线"],
        [
            ["", "", "", "", "", "是 / 否"],
            ["", "", "", "", "", "是 / 否"],
            ["", "", "", "", "", "是 / 否"],
            ["", "", "", "", "", "是 / 否"],
            ["", "", "", "", "", "是 / 否"],
        ],
        [1.8, 1.8, 3.7, 3.2, 3.2, 2.4],
        font_size=8.2,
    )

    add_note_table(
        doc,
        "上线判定规则",
        "只要最低上线门槛中任一项未通过，原则上不得提交正式发布。若客户要求有条件上线，必须由吴彤彤在《最终上线确认单》中写明风险、兜底方案、负责人和预计修复时间。",
        RISK_FILL,
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
